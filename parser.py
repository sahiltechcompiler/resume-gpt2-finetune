import os
import pdfplumber

import logging
from docx import Document
from PIL import Image
import pytesseract


from transformers import GPT2LMHeadModel, GPT2Tokenizer
import torch

# Load the fine-tuned model
model_path = "gpt2_resume_parser_model"
tokenizer = GPT2Tokenizer.from_pretrained(model_path)
model = GPT2LMHeadModel.from_pretrained(model_path)
model.eval()
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

def call_gpt2_resume_parser(resume_text):
    prompt = f"""<|startoftext|>
Input: {resume_text}
Output:"""
    inputs = tokenizer(prompt, return_tensors="pt", padding=True, truncation=True, max_length=1024)
    input_ids = inputs["input_ids"].to(device)
    attention_mask = inputs["attention_mask"].to(device)

    outputs = model.generate(
        input_ids,
        attention_mask=attention_mask,
        max_length=1024,
        do_sample=True,
        top_p=0.95,
        top_k=50,
        temperature=0.7,
        pad_token_id=tokenizer.eos_token_id,
    )

    result = tokenizer.decode(outputs[0], skip_special_tokens=True)
    if "Output:" in result:
        json_output = result.split("Output:", 1)[1].strip()
        print(json_output)
        return json_output
    else:
        print("⚠️ Output section not found in model response.")
        return ""
logging.getLogger("pdfminer").setLevel(logging.ERROR)

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "llama3.2"
RESUME_FOLDER = "resumes"


def extract_text_from_pdf(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
    except Exception as e:
        print(f"⚠️ Error extracting text from {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception as e:
        print(f"⚠️ Error extracting text from {file_path}: {e}")
        return ""


def extract_text_from_doc(file_path):
    try:
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception:
        print(f"🔄 Converting .doc to .docx: {file_path}")
        try:
            output_dir = os.path.dirname(file_path)
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            converted_path = os.path.join(output_dir, base_name + ".docx")

            os.system(f'soffice --headless --convert-to docx "{file_path}" --outdir "{output_dir}"')

            if os.path.exists(converted_path):
                doc = Document(converted_path)
                text = "\n".join(p.text for p in doc.paragraphs).strip()
                os.remove(converted_path)
                print(f"🗑️ Deleted converted .docx file: {converted_path}")
                return text
        except Exception as e:
            print(f"⚠️ Failed to convert .doc to .docx: {e}")
    return ""


def extract_text_from_image(file_path):
    try:
        image = Image.open(file_path)
        return pytesseract.image_to_string(image).strip()
    except Exception as e:
        print(f"⚠️ Error extracting text from image {file_path}: {e}")
        return ""






def main():
    
    for file_name in os.listdir(RESUME_FOLDER):
        file_path = os.path.join(RESUME_FOLDER, file_name)
        text = ""

        if file_name.lower().endswith(".pdf"):
            print(f"📄 Parsing PDF: {file_name}")
            text = extract_text_from_pdf(file_path)
        elif file_name.lower().endswith(".docx"):
            print(f"📝 Parsing DOCX: {file_name}")
            text = extract_text_from_docx(file_path)
        elif file_name.lower().endswith(".doc"):
            print(f"📝 Parsing DOC (old format): {file_name}")
            text = extract_text_from_doc(file_path)
        elif file_name.lower().endswith((".png", ".jpg", ".jpeg")):
            print(f"🖼️ Parsing Image: {file_name}")
            text = extract_text_from_image(file_path)
        else:
            print(f"⛔ Unsupported file type: {file_name}. Skipping.")
            continue

        if not text:
            print(f"⚠️ No text extracted from {file_name}. Skipping.")
            continue

        call_gpt2_resume_parser(text)
        

    


if __name__ == "__main__":
    main()